from docx import Document
from flask import Flask, request, send_file, render_template_string
import smtplib
import sqlite3
from email.message import EmailMessage
import matplotlib.pyplot as plt
import os

# Caminhos de salvamento
db_path = '/mnt/data/contratos.db'
chart_path = '/mnt/data/lucro_por_servico.png'

# Criar banco de dados
conn = sqlite3.connect(db_path)
cursor = conn.cursor()
cursor.execute('''CREATE TABLE IF NOT EXISTS contratos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    nome TEXT,
                    email TEXT,
                    servico TEXT,
                    valor REAL)''')
conn.commit()
conn.close()

# Criar gráfico de projeção financeira
servicos = ['Análise de Dados', 'Consultoria', 'Perícia', 'Projetos']
valores = [1500000, 1200000, 900000, 1800000]
custos = [500000, 400000, 300000, 600000]
lucros = [v - c for v, c in zip(valores, custos)]

plt.figure(figsize=(8, 6))
plt.bar(servicos, lucros)
plt.title('Projeção de Lucro por Serviço')
plt.xlabel('Serviço')
plt.ylabel('Lucro (USD)')
plt.tight_layout()
plt.savefig(chart_path)
plt.close()

# Função para gerar contrato em DOCX
def gerar_contrato(nome, servico, valor):
    doc = Document()
    doc.add_heading('Contrato de Prestação de Serviço', 0)
    doc.add_paragraph(f'Cliente: {nome}')
    doc.add_paragraph(f'Serviço Contratado: {servico}')
    doc.add_paragraph(f'Valor: USD {valor:.2f}')
    path = f'/mnt/data/contrato_{nome.replace(" ", "_")}.docx'
    doc.save(path)
    return path

# Interface Flask
app = Flask(__name__)

html_form = """
<!DOCTYPE html>
<html lang="pt-br">
<head><meta charset="UTF-8"><title>Cadastro de Contrato</title></head>
<body>
  <h1>Cadastro de Contrato</h1>
  <form action="/gerar" method="post">
    Nome: <input type="text" name="nome"><br>
    E-mail: <input type="email" name="email"><br>
    Serviço: <input type="text" name="servico"><br>
    Valor: <input type="number" name="valor"><br>
    <input type="submit" value="Gerar e Simular Envio">
  </form>
  <h2>Projeção de Lucro por Serviço</h2>
  <img src="/static/lucro_por_servico.png" width="600px">
</body>
</html>
"""

@app.route('/')
def index():
    return render_template_string(html_form)

@app.route('/gerar', methods=['POST'])
def gerar():
    nome = request.form['nome']
    email = request.form['email']
    servico = request.form['servico']
    valor = float(request.form['valor'])

    # Salvar no banco
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("INSERT INTO contratos (nome, email, servico, valor) VALUES (?, ?, ?, ?)",
                   (nome, email, servico, valor))
    conn.commit()
    conn.close()

    # Gerar contrato
    contrato_path = gerar_contrato(nome, servico, valor)

    # Simular envio de e-mail (apenas log)
    print(f"Simulado envio de {contrato_path} para {email}")

    return f'Contrato gerado e enviado para {email} com sucesso!'

# Lista de arquivos gerados
files = [db_path, chart_path]
files.extend([f'/mnt/data/{f}' for f in os.listdir('/mnt/data') if f.startswith('contrato_')])

files
